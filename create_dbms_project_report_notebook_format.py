from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "DBMS_Project_Report_Notebook_Format.docx"
SCHEMA_IMG = ROOT / "database_schema_diagram.png"
ER_IMG = ROOT / "er_diagram.png"

PROJECT_TITLE = "AI-Powered Placement Preparation and Student Progress Tracker"
HEADER_TITLE = "Placement Preparation and Student Progress Tracker"
FOOTER_LEFT = "Dept. of AIML / BMIT"
FOOTER_RIGHT = "2025-26"

ACCENT = RGBColor(16, 78, 105)
LIGHT_FILL = "EAF3F7"


entities = {
    "User": ["_id", "name", "email", "password", "role", "isActive", "lastLoginAt"],
    "StudentProfile": ["_id", "userId", "mentorId", "department", "semester", "targetRole", "skills", "parsedResume", "readinessScore"],
    "Task": ["_id", "userId", "studentId", "assignedBy", "title", "category", "deadline", "priority", "status"],
    "ProgressLog": ["_id", "userId", "studentId", "date", "topicsCompleted", "studyMinutes", "mockInterviewScore", "resumeScore"],
    "MockInterview": ["_id", "userId", "studentId", "type", "company", "questions", "answers", "overallScore", "status"],
    "ATSReview": ["_id", "userId", "studentId", "resumeText", "jobDescription", "matchedKeywords", "missingKeywords", "metrics"],
    "PlacementApplication": ["_id", "userId", "studentId", "company", "role", "status", "rounds", "timeline"],
    "CodingProblem": ["_id", "slug", "title", "difficulty", "category", "companies", "testCases"],
    "CodingSubmission": ["_id", "userId", "studentId", "problemId", "language", "code", "verdict", "score"],
    "Conversation": ["_id", "participants", "lastMessage", "lastMessageAt", "unreadCount"],
    "Message": ["_id", "conversationId", "senderId", "receiverId", "content", "messageType", "isRead"],
    "Notification": ["_id", "userId", "title", "message", "type", "priority", "status"],
    "MentorNote": ["_id", "mentorId", "studentId", "visibility", "tone", "note", "actionItems"],
    "Analytics": ["_id", "scope", "userId", "studentId", "metrics", "status"],
}


def image_font(size=12):
    try:
        return ImageFont.truetype("times.ttf", size)
    except Exception:
        return ImageFont.load_default()


def draw_box(draw, x, y, w, h, title, rows, fill="#FFFFFF", outline="#426B7A"):
    draw.rounded_rectangle((x, y, x + w, y + h), radius=12, fill=fill, outline=outline, width=2)
    draw.rounded_rectangle((x, y, x + w, y + 34), radius=12, fill="#DDEFF5", outline=outline, width=2)
    draw.text((x + 10, y + 8), title, fill="#103447", font=image_font(15))
    row_y = y + 42
    for row in rows[:8]:
        draw.text((x + 10, row_y), row, fill="#1F2937", font=image_font(12))
        row_y += 18


def line_with_arrow(draw, start, end, label=""):
    draw.line((start, end), fill="#526D7A", width=3)
    x1, y1 = start
    x2, y2 = end
    pts = [(x2, y2), (x2 - 10 if x2 >= x1 else x2 + 10, y2 - 6), (x2 - 10 if x2 >= x1 else x2 + 10, y2 + 6)]
    draw.polygon(pts, fill="#526D7A")
    if label:
        lx = (x1 + x2) // 2
        ly = (y1 + y2) // 2 - 16
        draw.rectangle((lx - 42, ly - 2, lx + 42, ly + 18), fill="#F8FBFC")
        draw.text((lx - 34, ly), label, fill="#23343B", font=image_font(11))


def create_schema_diagram():
    img = Image.new("RGB", (1800, 1300), "#F8FBFC")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "Database Schema Diagram", fill="#0B2F43", font=image_font(32))
    draw.text((60, 78), PROJECT_TITLE, fill="#415A66", font=image_font(18))
    positions = {
        "User": (70, 150), "StudentProfile": (540, 150), "Task": (1010, 150), "ProgressLog": (1390, 150),
        "MockInterview": (70, 460), "ATSReview": (540, 460), "PlacementApplication": (1010, 460), "Analytics": (1390, 460),
        "CodingProblem": (70, 770), "CodingSubmission": (540, 770), "Conversation": (1010, 770), "Message": (1390, 770),
        "Notification": (300, 1070), "MentorNote": (780, 1070)
    }
    w, h = 330, 230
    for name, pos in positions.items():
        draw_box(draw, pos[0], pos[1], w, h, name, entities[name])
    for left, right in [
        ("User", "StudentProfile"), ("StudentProfile", "Task"), ("StudentProfile", "ProgressLog"),
        ("StudentProfile", "MockInterview"), ("StudentProfile", "ATSReview"), ("StudentProfile", "PlacementApplication"),
        ("StudentProfile", "Analytics"), ("CodingProblem", "CodingSubmission"), ("Conversation", "Message"),
        ("User", "Notification"), ("User", "MentorNote")
    ]:
        x1, y1 = positions[left]
        x2, y2 = positions[right]
        line_with_arrow(draw, (x1 + w, y1 + h // 2), (x2, y2 + h // 2))
    img.save(SCHEMA_IMG)


def create_er_diagram():
    img = Image.new("RGB", (1800, 1350), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 40), "ER Diagram", fill="#0B2F43", font=image_font(34))
    draw.text((70, 84), "Logical relationships among the main database entities", fill="#415A66", font=image_font(18))
    pos = {
        "User": (720, 150), "StudentProfile": (720, 430),
        "Task": (110, 370), "ProgressLog": (110, 630), "MockInterview": (110, 890),
        "ATSReview": (720, 760), "PlacementApplication": (1240, 370), "MentorNote": (1240, 630),
        "CodingProblem": (720, 1080), "CodingSubmission": (1240, 1010),
        "Conversation": (1240, 150), "Message": (1240, 890), "Notification": (110, 150), "Analytics": (720, 1010)
    }
    w, h = 360, 120
    for name, (x, y) in pos.items():
        rows = ["PK: _id"] + [f"FK: {r}" for r in entities[name] if r.endswith("Id") or r == "participants"][:3]
        draw_box(draw, x, y, w, h, name, rows, fill="#FCFEFF")
    for a, b, lbl in [
        ("User", "StudentProfile", "1 : 1"), ("StudentProfile", "Task", "1 : M"),
        ("StudentProfile", "ProgressLog", "1 : M"), ("StudentProfile", "MockInterview", "1 : M"),
        ("StudentProfile", "ATSReview", "1 : M"), ("StudentProfile", "PlacementApplication", "1 : M"),
        ("StudentProfile", "MentorNote", "1 : M"), ("StudentProfile", "Analytics", "1 : M"),
        ("CodingProblem", "CodingSubmission", "1 : M"), ("User", "Notification", "1 : M"),
        ("Conversation", "Message", "1 : M")
    ]:
        ax, ay = pos[a]
        bx, by = pos[b]
        start = (ax + (0 if bx < ax else w), ay + h // 2)
        end = (bx + (w if bx < ax else 0), by + h // 2)
        line_with_arrow(draw, start, end, lbl)
    img.save(ER_IMG)


def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_font(run, 10)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def clear_header_footer(section):
    for part in [section.header, section.footer]:
        for p in part.paragraphs:
            p.text = ""


def apply_main_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(HEADER_TITLE)
    set_font(r, 10)
    footer = section.footer.paragraphs[0]
    footer.text = ""
    table = section.footer.add_table(rows=1, cols=3, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2.5, 1.5, 2.5]
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = Inches(width)
    set_cell_text(table.cell(0, 0), FOOTER_LEFT, size=10)
    center_p = table.cell(0, 1).paragraphs[0]
    center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_p.text = ""
    add_page_field(center_p)
    set_cell_text(table.cell(0, 2), FOOTER_RIGHT, size=10)
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                edge = OxmlElement(f"w:{side}")
                edge.set(qn("w:val"), "nil")
                borders.append(edge)
            tc_pr.append(borders)


def apply_manual_front_footer(section, roman):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_header_footer(section)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(roman)
    set_font(r, 10)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    clear_header_footer(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for style_name, size, align in [
        ("Heading 1", 16, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 12, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        st = styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.paragraph_format.alignment = align
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def new_page_section(doc, footer=None, main=False):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    if main:
        apply_main_header_footer(sec)
    elif footer:
        apply_manual_front_footer(sec, footer)
    else:
        clear_header_footer(sec)
    return sec


def heading(doc, text, level=1):
    p = doc.add_heading(text.upper() if level == 1 else text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return p


def para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_font(r, 12)


def add_table(doc, caption, headers, rows):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, 12, True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT_TITLE.upper())
    set_font(r, 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DBMS PROJECT REPORT")
    set_font(r, 16, True)
    for _ in range(2):
        doc.add_paragraph()
    rows = [
        ("Submitted By", "____________________________"),
        ("Register Number", "____________________________"),
        ("Department", "AIML"),
        ("Semester", "IV"),
        ("Academic Year", "2025-26"),
        ("Institution", "BMIT"),
    ]
    add_table(doc, "Project Details", ["Particulars", "Details"], rows)
    para(doc, "This report describes the database design, schema, ER model, implementation, and result analysis of the Placement Preparation and Student Progress Tracker project.")


def add_front_matter(doc):
    new_page_section(doc, "i")
    heading(doc, "Abstract")
    for text in [
        "The project AI-Powered Placement Preparation and Student Progress Tracker is a DBMS-based web application developed to organize student placement preparation activities in one system. The application stores student profiles, resume details, tasks, progress logs, mock interviews, coding submissions, mentor notes, placement applications, messages, notifications, and analytics records.",
        "The database is implemented using MongoDB with Mongoose schemas. MongoDB is suitable for this project because the data contains both structured fields and nested objects such as interview answers, resume versions, application rounds, timelines, and coding testcase results. The backend is developed with Node.js and Express.js, while the frontend is developed using React.",
        "The main objective of the project is to create a connected preparation platform where students can track their placement readiness and mentors can guide students using stored progress data. The report explains the database schema, ER diagram, hardware and software requirements, database design, implementation details, results, conclusion, and references.",
    ]:
        para(doc, text)
    new_page_section(doc, "ii")
    heading(doc, "Table of Contents")
    add_table(doc, "Table 1.1: Table of Contents", ["Heading", "Page No."], [
        ("Abstract", "i"), ("Table of Contents", "ii"), ("Table of Figures", "iii"), ("List of Tables", "iv"),
        ("Introduction", "1"), ("Database Schema Diagram", "4"), ("ER Diagram", "7"),
        ("Hardware and Software Requirement", "9"), ("Database Design and Implementation", "11"),
        ("Results", "17"), ("Conclusion", "19"), ("References", "20"),
    ])
    new_page_section(doc, "iii")
    heading(doc, "Table of Figures")
    add_table(doc, "Table 1.2: Table of Figures", ["Figure No.", "Figure Name", "Page No."], [
        ("Figure 1.1", "Database Schema Diagram", "5"),
        ("Figure 1.2", "ER Diagram", "8"),
    ])
    new_page_section(doc, "iv")
    heading(doc, "List of Tables")
    add_table(doc, "Table 1.3: List of Tables", ["Table No.", "Table Name", "Page No."], [
        ("Table 1.1", "Table of Contents", "ii"), ("Table 1.2", "Table of Figures", "iii"),
        ("Table 1.3", "List of Tables", "iv"), ("Table 1.4", "Main Collections and Purpose", "6"),
        ("Table 1.5", "Relationship Description", "8"), ("Table 1.6", "Hardware Requirements", "9"),
        ("Table 1.7", "Software Requirements", "10"), ("Table 1.8", "Model Implementation Details", "12"),
        ("Table 1.9", "CRUD Operations", "15"), ("Table 1.10", "Result Summary", "18"),
    ])


def add_main_content(doc):
    new_page_section(doc, main=True)
    heading(doc, "Introduction")
    for text in [
        "The Placement Preparation and Student Progress Tracker is developed as a database-oriented application for managing the activities involved in campus placement preparation. In a normal college environment, students prepare for aptitude tests, coding rounds, resume screening, technical interviews, and HR rounds using separate files or platforms. This often makes it difficult to know the actual preparation status of a student.",
        "This project brings those activities into one structured system. A student can maintain profile details, upload resume information, list target companies, record weak topics, solve coding problems, attend mock interviews, track applications, and view readiness scores. A mentor can monitor the same student through tasks, notes, feedback, activity logs, and progress information.",
        "The project uses the MERN stack. MongoDB is used as the database, Express.js and Node.js are used for the backend API, and React is used for the frontend. Mongoose is used to define database schemas, references, indexes, validation rules, and embedded subdocuments.",
        "From the DBMS point of view, the project is useful because it contains multiple related collections. The main database objects are User, StudentProfile, Task, ProgressLog, MockInterview, ATSReview, PlacementApplication, CodingProblem, CodingSubmission, Conversation, Message, Notification, MentorNote, and Analytics.",
    ]:
        para(doc, text)
    heading(doc, "Objectives", 2)
    for item in [
        "To create a database-backed platform for placement preparation.",
        "To store student academic details, resume data, tasks, progress, interviews, and applications.",
        "To connect related records using ObjectId references in MongoDB.",
        "To support dashboard analytics based on stored preparation activity.",
        "To provide role-based access for students, mentors, and administrators.",
    ]:
        bullet(doc, item)
    doc.add_page_break()
    heading(doc, "Need for the System", 2)
    for text in [
        "Students usually maintain preparation data in different places. Resume files may be stored separately, tasks may be written in notebooks, coding practice may happen on another platform, and mentor feedback may remain in chat messages. Because of this separation, it becomes difficult to measure preparation improvement over time.",
        "The proposed system solves this problem by storing related preparation data in one database. The system not only stores basic student information but also stores the history of preparation activities. This includes progress logs, interview scores, resume versions, coding submissions, application status, and mentor notes.",
        "The database design reduces duplication by using references. For example, mentor details are not copied into every task or note. The system stores mentorId as a reference. Similarly, student-owned records point to StudentProfile through studentId.",
    ]:
        para(doc, text)
    add_table(doc, "Table 1.4: User Roles and Responsibilities", ["User Type", "Role in the System"], [
        ("Student", "Maintains profile, uploads resume, attempts interviews, solves coding problems, records progress, and tracks applications."),
        ("Mentor", "Assigns tasks, reviews progress, writes notes, gives feedback, and monitors student readiness."),
        ("Admin", "Observes broader analytics and manages system-level information."),
    ])
    doc.add_page_break()

    heading(doc, "Database Schema Diagram")
    para(doc, "The database schema diagram shows the main MongoDB collections used in the project. Each collection stores one major part of the application. High-volume activity records such as interviews, messages, progress logs, coding submissions, and applications are kept as separate collections.")
    doc.add_picture(str(SCHEMA_IMG), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap.add_run("Figure 1.1: Database Schema Diagram"), 12, True)
    para(doc, "The User and StudentProfile collections are the central collections. The User collection stores account and authentication details. StudentProfile stores academic and placement information. Other collections use userId and studentId fields to maintain proper links with the student account and profile.")
    doc.add_page_break()
    add_table(doc, "Table 1.5: Main Collections and Purpose", ["Collection", "Purpose"], [
        ("User", "Stores login identity, encrypted password, role, and account status."),
        ("StudentProfile", "Stores academic details, skills, resume data, readiness scores, and mentor mapping."),
        ("Task", "Stores preparation tasks with status, category, priority, and deadline."),
        ("ProgressLog", "Stores study minutes, completed topics, scores, and notes."),
        ("MockInterview", "Stores questions, answers, scores, feedback, and delivery metrics."),
        ("ATSReview", "Stores resume review output and keyword matching metrics."),
        ("PlacementApplication", "Stores company applications, rounds, timeline, and status."),
        ("CodingProblem", "Stores coding challenges, constraints, templates, and testcases."),
        ("CodingSubmission", "Stores submitted code, verdict, score, and testcase results."),
    ])
    para(doc, "The schema uses embedded arrays for data that belongs fully to a parent document. Examples include interview answers, resume versions, application rounds, timeline events, coding testcases, and submission cases. Separate collections are used when the records can grow independently or need separate querying.")
    doc.add_page_break()

    heading(doc, "ER Diagram")
    para(doc, "The ER diagram explains the logical relationship between the entities. Although MongoDB is a document database, ER modeling is still useful for understanding the database design before implementation.")
    doc.add_picture(str(ER_IMG), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap.add_run("Figure 1.2: ER Diagram"), 12, True)
    doc.add_page_break()
    add_table(doc, "Table 1.6: Relationship Description", ["Relationship", "Cardinality", "Explanation"], [
        ("User - StudentProfile", "1 : 1", "A student account has one detailed placement profile."),
        ("StudentProfile - Task", "1 : M", "A student can have many preparation tasks."),
        ("StudentProfile - ProgressLog", "1 : M", "A student can record many progress entries."),
        ("StudentProfile - MockInterview", "1 : M", "A student can attend many mock interviews."),
        ("StudentProfile - ATSReview", "1 : M", "A student can have many resume reviews."),
        ("StudentProfile - PlacementApplication", "1 : M", "A student can apply to many companies."),
        ("CodingProblem - CodingSubmission", "1 : M", "One problem can have many submissions."),
        ("Conversation - Message", "1 : M", "A conversation contains many messages."),
    ])
    para(doc, "The one-to-many relationships are important because preparation is a continuous activity. A student does not have only one task or one mock interview. The database must preserve older records so that progress can be measured later.")
    doc.add_page_break()

    heading(doc, "Hardware and Software Requirement")
    para(doc, "The project can be developed and tested on a normal laptop. Since the database is hosted in MongoDB Atlas, a stable internet connection is required during development and testing.")
    add_table(doc, "Table 1.7: Hardware Requirements", ["Component", "Minimum Requirement", "Recommended Requirement"], [
        ("Processor", "Dual-core processor", "Intel i5 / Ryzen 5 or above"),
        ("RAM", "4 GB", "8 GB or above"),
        ("Storage", "2 GB free space", "5 GB or above"),
        ("Network", "Internet connection", "Stable broadband connection"),
        ("Display", "1366 x 768", "Full HD display"),
    ])
    doc.add_page_break()
    add_table(doc, "Table 1.8: Software Requirements", ["Software", "Purpose"], [
        ("Node.js", "Runs backend server and frontend tooling."),
        ("MongoDB Atlas", "Stores project database in cloud."),
        ("Express.js", "Provides REST API routes."),
        ("React.js", "Builds frontend user interface."),
        ("Mongoose", "Defines schemas, validation, references, and indexes."),
        ("Axios", "Connects frontend with backend API."),
        ("JWT and bcryptjs", "Used for authentication and password hashing."),
        ("Visual Studio Code", "Used for project development."),
        ("Web Browser", "Used to access and test the application."),
    ])
    heading(doc, "Functional Requirements", 2)
    for item in [
        "Users must be able to register and log in securely.",
        "Students must be able to create and update their profile.",
        "The system must store resume data and resume review history.",
        "Students must be able to create tasks and record progress.",
        "Mock interview questions, answers, scores, and feedback must be saved.",
        "Placement applications must be tracked with status and rounds.",
    ]:
        bullet(doc, item)
    doc.add_page_break()

    heading(doc, "Database Design and Implementation")
    for text in [
        "Database design is the main part of this project. The design was prepared by identifying the users of the system, the data generated by each user, and the relationships between the data. MongoDB is used because the project contains both fixed fields and nested data.",
        "Mongoose models are used for implementation. Each schema defines the field names, data types, required conditions, enum values, default values, references, and indexes. Controllers use these models to perform create, read, update, and delete operations.",
    ]:
        para(doc, text)
    add_table(doc, "Table 1.9: Model Implementation Details", ["Model", "Implementation Detail"], [
        ("User", "Stores login data, role, active status, and hashed password."),
        ("StudentProfile", "Stores placement profile, resume metadata, skills, readiness scores, and mentor mapping."),
        ("Task", "Stores title, category, deadline, priority, status, source, and tags."),
        ("MockInterview", "Stores interview questions, answers, score, feedback, and metrics."),
        ("ATSReview", "Stores resume text, job description, keyword matching, and ATS metrics."),
        ("PlacementApplication", "Stores company, role, application status, rounds, and timeline."),
        ("CodingProblem", "Stores problem statement, difficulty, tags, templates, and testcases."),
        ("CodingSubmission", "Stores user code, verdict, score, runtime, and testcase output."),
    ])
    doc.add_page_break()
    heading(doc, "Normalization and Embedding", 2)
    for text in [
        "The database uses a mixed design. Separate collections are used for independent records, while embedded subdocuments are used for data that belongs strongly to one parent document. This approach is suitable for MongoDB.",
        "For example, application rounds are embedded inside PlacementApplication because those rounds are meaningful only for that application. Interview answers are embedded inside MockInterview because they belong to one interview session. Messages are stored separately because a conversation can contain many messages and may need pagination.",
        "References are used when records must be connected across collections. The commonly used reference fields are userId, studentId, mentorId, problemId, conversationId, senderId, and receiverId.",
    ]:
        para(doc, text)
    heading(doc, "Indexing", 2)
    para(doc, "Indexes are used on fields that are frequently used for searching, filtering, and sorting. Examples include email, role, userId, studentId, status, priority, createdAt, company, difficulty, category, slug, conversationId, and lastMessageAt. These indexes help the dashboard and history pages load records faster.")
    doc.add_page_break()
    add_table(doc, "Table 1.10: CRUD Operations", ["Operation", "Example in Project"], [
        ("Create", "Register user, create task, upload resume review, submit coding solution."),
        ("Read", "Load dashboard, fetch profile, list tasks, view ATS history."),
        ("Update", "Update profile, mark task completed, change application status."),
        ("Delete / Soft Delete", "Messages and status-based records can be hidden or marked inactive."),
    ])
    heading(doc, "Authentication and Security", 2)
    para(doc, "The system uses JWT for authentication. Passwords are hashed with bcrypt before saving in the database. Protected routes check the token and identify the logged-in user. Role middleware ensures that students, mentors, and admins access only their permitted routes.")
    heading(doc, "Backend API Flow", 2)
    para(doc, "The frontend sends requests to the Express API. The API verifies authentication, validates inputs, performs database operations using Mongoose models, and sends JSON responses back to the frontend. The frontend never connects directly to MongoDB, which improves security.")
    doc.add_page_break()
    heading(doc, "Data Flow Example", 2)
    for text in [
        "When a student updates the profile, the frontend sends changed fields to the profile API. The backend verifies the token and updates the StudentProfile collection.",
        "When a resume is uploaded, the backend extracts text, parses skills and projects, updates parsedResume, stores resume metadata, and keeps a resume version history.",
        "When a mock interview is completed, questions, answers, evaluation metrics, and overall score are saved in MockInterview. This record can later be used for analytics.",
        "When a task is completed, the task status and completedAt fields are updated. This improves task completion tracking and mentor review.",
    ]:
        para(doc, text)
    doc.add_page_break()

    heading(doc, "Results")
    para(doc, "The completed project provides a working database-backed placement preparation system. The database stores student information, preparation tasks, resume data, interview history, coding attempts, mentor guidance, notifications, messages, and application records.")
    add_table(doc, "Table 1.11: Result Summary", ["Feature", "Collection Used", "Result"], [
        ("Profile Management", "StudentProfile", "Student placement metadata is stored and retrieved."),
        ("Resume Analysis", "StudentProfile, ATSReview", "Resume text, skills, and review history are saved."),
        ("Mock Interview", "MockInterview", "Questions, answers, scores, and feedback are stored."),
        ("Task Planning", "Task", "Tasks can be filtered by status, priority, and category."),
        ("Coding Practice", "CodingProblem, CodingSubmission", "Problems and submissions are linked by problemId."),
        ("Application Tracker", "PlacementApplication", "Company application status and rounds are maintained."),
        ("Communication", "Conversation, Message", "Messages are grouped under conversations."),
    ])
    doc.add_page_break()
    heading(doc, "Testing Observations", 2)
    for item in [
        "Profile workflow confirms that student details can be created, updated, and fetched.",
        "Task workflow confirms that status changes and filtering work correctly.",
        "Resume workflow confirms that resume metadata and parsed information can be stored.",
        "Interview workflow confirms that question-wise answer and score storage is possible.",
        "Application workflow confirms that company-wise status and timeline can be maintained.",
    ]:
        bullet(doc, item)
    para(doc, "The result shows that the database design supports both current status and historical tracking. This is important because placement preparation improves over time and older records should not be lost.")
    doc.add_page_break()

    heading(doc, "Conclusion")
    for text in [
        "The Placement Preparation and Student Progress Tracker successfully demonstrates the use of DBMS concepts in a practical web application. The project contains multiple related collections, schema validation, indexes, embedded subdocuments, references, authentication, and role-based access.",
        "MongoDB is suitable for this project because the data contains flexible structures such as resume versions, interview answers, application timelines, coding testcases, and notification metadata. Mongoose gives structure to this flexible data through schemas and validation rules.",
        "The project helps students maintain preparation data in one place and helps mentors monitor student readiness. The database design can also be extended in the future for more analytics, company-wise roadmaps, advanced reports, and improved mentor dashboards.",
    ]:
        para(doc, text)
    heading(doc, "Future Enhancements", 2)
    for item in [
        "Add more company-specific preparation plans.",
        "Generate detailed PDF reports for students and mentors.",
        "Add stronger analytics using historical progress trends.",
        "Improve notification scheduling and reminders.",
        "Add more backup and audit features for production use.",
    ]:
        bullet(doc, item)
    doc.add_page_break()

    heading(doc, "References")
    refs = [
        "MongoDB Documentation, MongoDB Manual: Documents, Collections, Indexes, and Schema Design.",
        "Mongoose Documentation, Schemas, Models, Validation, Middleware, and Query Indexing.",
        "Express.js Documentation, Routing, Middleware, Error Handling, and REST API Development.",
        "React Documentation, Components, State Management, and Frontend Rendering.",
        "Node.js Documentation, JavaScript Runtime and Server-side Application Development.",
        "JSON Web Token Documentation, Token-based Authentication Concepts.",
        "bcryptjs Package Documentation, Password Hashing for Web Applications.",
        "Project Source Code, AI-Powered Placement Preparation and Student Progress Tracker.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref)
    doc.add_page_break()
    heading(doc, "Project Files Referred", 2)
    add_table(doc, "Table 1.12: Project Files Referred", ["Project File", "Reason for Reference"], [
        ("backend/models/User.js", "Authentication entity and user role details."),
        ("backend/models/StudentProfile.js", "Student profile, resume, skills, and readiness schema."),
        ("backend/models/Task.js", "Preparation task schema and status fields."),
        ("backend/models/MockInterview.js", "Interview question, answer, and score storage."),
        ("backend/models/ATSReview.js", "Resume review and ATS score schema."),
        ("backend/models/PlacementApplication.js", "Application tracking and round details."),
        ("backend/models/CodingProblem.js", "Coding problem and testcase structure."),
        ("backend/models/CodingSubmission.js", "Coding submission and verdict details."),
        ("README.md", "Project overview and technology stack."),
    ])


def build():
    create_schema_diagram()
    create_er_diagram()
    doc = Document()
    style_doc(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_main_content(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
