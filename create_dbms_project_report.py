from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "DBMS_Project_Report_Placement_Preparation_Tracker.docx"
SCHEMA_IMG = ROOT / "database_schema_diagram.png"
ER_IMG = ROOT / "er_diagram.png"


ACCENT = RGBColor(21, 96, 130)
DARK = RGBColor(31, 42, 55)
LIGHT_FILL = "EAF3F7"


entities = {
    "User": ["_id", "name", "email", "password", "role", "isActive", "lastLoginAt"],
    "StudentProfile": ["_id", "userId", "mentorId", "department", "semester", "targetRole", "skills", "parsedResume", "readinessScore"],
    "Task": ["_id", "userId", "studentId", "assignedBy", "title", "category", "deadline", "priority", "status"],
    "ProgressLog": ["_id", "userId", "studentId", "date", "topicsCompleted", "studyMinutes", "mockInterviewScore", "resumeScore"],
    "MockInterview": ["_id", "userId", "studentId", "type", "company", "questions", "answers", "overallScore", "status"],
    "ATSReview": ["_id", "userId", "studentId", "resumeText", "jobDescription", "matchedKeywords", "missingKeywords", "metrics"],
    "PlacementApplication": ["_id", "userId", "studentId", "company", "role", "status", "rounds", "timeline"],
    "CodingProblem": ["_id", "slug", "title", "difficulty", "category", "companies", "testCases"],
    "CodingSubmission": ["_id", "userId", "studentId", "problemId", "language", "code", "verdict", "score"],
    "Conversation": ["_id", "participants", "lastMessage", "lastMessageAt", "unreadCount"],
    "Message": ["_id", "conversationId", "senderId", "receiverId", "content", "messageType", "isRead"],
    "Notification": ["_id", "userId", "title", "message", "type", "priority", "status"],
    "MentorNote": ["_id", "mentorId", "studentId", "visibility", "tone", "note", "actionItems"],
    "Analytics": ["_id", "scope", "userId", "studentId", "metrics", "status"],
}


relationships = [
    ("User", "StudentProfile", "1", "1", "owns"),
    ("User", "Task", "1", "M", "creates/receives"),
    ("StudentProfile", "Task", "1", "M", "tracks"),
    ("StudentProfile", "ProgressLog", "1", "M", "records"),
    ("StudentProfile", "MockInterview", "1", "M", "attempts"),
    ("StudentProfile", "ATSReview", "1", "M", "reviews"),
    ("StudentProfile", "PlacementApplication", "1", "M", "applies"),
    ("CodingProblem", "CodingSubmission", "1", "M", "evaluates"),
    ("User", "CodingSubmission", "1", "M", "submits"),
    ("Conversation", "Message", "1", "M", "contains"),
    ("User", "Notification", "1", "M", "receives"),
    ("User", "MentorNote", "1", "M", "writes"),
    ("StudentProfile", "MentorNote", "1", "M", "gets"),
    ("StudentProfile", "Analytics", "1", "M", "summarized by"),
]


def font(size=12, bold=False):
    try:
        return ImageFont.truetype("times.ttf", size)
    except Exception:
        return ImageFont.load_default()


def draw_wrapped(draw, xy, text, max_chars, fill, fnt):
    x, y = xy
    for line in wrap(text, max_chars):
        draw.text((x, y), line, fill=fill, font=fnt)
        y += 18
    return y


def draw_box(draw, x, y, w, h, title, rows, fill="#FFFFFF", outline="#426B7A"):
    draw.rounded_rectangle((x, y, x + w, y + h), radius=12, fill=fill, outline=outline, width=2)
    draw.rounded_rectangle((x, y, x + w, y + 34), radius=12, fill="#DDEFF5", outline=outline, width=2)
    draw.text((x + 10, y + 8), title, fill="#103447", font=font(15))
    row_y = y + 42
    for row in rows[:8]:
        draw.text((x + 10, row_y), row, fill="#1F2937", font=font(12))
        row_y += 18


def line_with_arrow(draw, start, end, label=""):
    draw.line((start, end), fill="#526D7A", width=3)
    x1, y1 = start
    x2, y2 = end
    # Simple arrow head.
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 10, y2 - 6), (x2 - 10, y2 + 6)]
    else:
        pts = [(x2, y2), (x2 + 10, y2 - 6), (x2 + 10, y2 + 6)]
    draw.polygon(pts, fill="#526D7A")
    if label:
        lx = (x1 + x2) // 2
        ly = (y1 + y2) // 2 - 16
        draw.rectangle((lx - 44, ly - 2, lx + 44, ly + 18), fill="#F8FBFC")
        draw.text((lx - 38, ly), label, fill="#23343B", font=font(11))


def create_schema_diagram():
    img = Image.new("RGB", (1800, 1300), "#F8FBFC")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "Database Schema Diagram - Main MongoDB Collections", fill="#0B2F43", font=font(30))
    draw.text((60, 78), "Placement Preparation and Student Progress Tracker", fill="#415A66", font=font(18))
    positions = {
        "User": (70, 150), "StudentProfile": (540, 150), "Task": (1010, 150), "ProgressLog": (1390, 150),
        "MockInterview": (70, 460), "ATSReview": (540, 460), "PlacementApplication": (1010, 460), "Analytics": (1390, 460),
        "CodingProblem": (70, 770), "CodingSubmission": (540, 770), "Conversation": (1010, 770), "Message": (1390, 770),
        "Notification": (300, 1070), "MentorNote": (780, 1070)
    }
    w, h = 330, 230
    for name, pos in positions.items():
        draw_box(draw, pos[0], pos[1], w, h, name, entities[name])
    schema_lines = [
        ("User", "StudentProfile"), ("StudentProfile", "Task"), ("StudentProfile", "ProgressLog"),
        ("StudentProfile", "MockInterview"), ("StudentProfile", "ATSReview"), ("StudentProfile", "PlacementApplication"),
        ("StudentProfile", "Analytics"), ("CodingProblem", "CodingSubmission"), ("Conversation", "Message"),
        ("User", "Notification"), ("User", "MentorNote")
    ]
    for left, right in schema_lines:
        x1, y1 = positions[left]
        x2, y2 = positions[right]
        line_with_arrow(draw, (x1 + w, y1 + h // 2), (x2, y2 + h // 2))
    img.save(SCHEMA_IMG)


def create_er_diagram():
    img = Image.new("RGB", (1800, 1350), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((70, 40), "ER Diagram - Logical View", fill="#0B2F43", font=font(32))
    draw.text((70, 84), "Relationships among users, student profiles, preparation records, and communication modules", fill="#415A66", font=font(18))
    pos = {
        "User": (720, 150), "StudentProfile": (720, 430),
        "Task": (110, 370), "ProgressLog": (110, 630), "MockInterview": (110, 890),
        "ATSReview": (720, 760), "PlacementApplication": (1240, 370), "MentorNote": (1240, 630),
        "CodingProblem": (720, 1080), "CodingSubmission": (1240, 1010),
        "Conversation": (1240, 150), "Message": (1240, 890), "Notification": (110, 150), "Analytics": (720, 1010)
    }
    w, h = 360, 120
    for name, (x, y) in pos.items():
        rows = [f"PK: _id"] + [f"FK: {r}" for r in entities[name] if r.endswith("Id") or r in ["participants"]][:3]
        draw_box(draw, x, y, w, h, name, rows, fill="#FCFEFF")
    rels = [
        ("User", "StudentProfile", "1 : 1"), ("StudentProfile", "Task", "1 : M"), ("StudentProfile", "ProgressLog", "1 : M"),
        ("StudentProfile", "MockInterview", "1 : M"), ("StudentProfile", "ATSReview", "1 : M"), ("StudentProfile", "PlacementApplication", "1 : M"),
        ("StudentProfile", "MentorNote", "1 : M"), ("StudentProfile", "Analytics", "1 : M"), ("CodingProblem", "CodingSubmission", "1 : M"),
        ("User", "Notification", "1 : M"), ("Conversation", "Message", "1 : M")
    ]
    for a, b, lbl in rels:
        ax, ay = pos[a]
        bx, by = pos[b]
        start = (ax + (0 if bx < ax else w), ay + h // 2)
        end = (bx + (w if bx < ax else 0), by + h // 2)
        line_with_arrow(draw, start, end, lbl)
    img.save(ER_IMG)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)
    header = section.header.paragraphs[0]
    header.text = "DBMS Project Report - Placement Preparation and Student Progress Tracker"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Times New Roman"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(90, 104, 114)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        st = styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT


def add_title(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Powered Placement Preparation and\nStudent Progress Tracker")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DBMS Project Report")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Times New Roman"
    for _ in range(3):
        doc.add_paragraph()
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Project Type", "Database Management System Project"),
        ("Technology Stack", "MongoDB, Express.js, React.js, Node.js"),
        ("Database", "MongoDB Atlas using Mongoose schema models"),
        ("Prepared By", "____________________________"),
        ("Register Number", "____________________________"),
        ("Department / Semester", "AIML / Semester 4"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        set_cell_text(row.cells[0], k, True)
        set_cell_text(row.cells[1], v)
        set_cell_shading(row.cells[0], LIGHT_FILL)
    doc.add_page_break()


def para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_FILL)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], str(value))
    doc.add_paragraph()
    return table


def add_toc(doc):
    heading(doc, "Contents")
    contents = [
        ("1. Introduction", "3"),
        ("2. Database-Schema Diagram", "6"),
        ("3. ER Diagram", "9"),
        ("4. Hardware and Software Requirement", "11"),
        ("5. Database Design and Implementation", "13"),
        ("6. Results", "18"),
        ("7. Conclusion", "20"),
        ("8. References", "21"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Section", True)
    set_cell_text(table.rows[0].cells[1], "Page No.", True)
    set_cell_shading(table.rows[0].cells[0], LIGHT_FILL)
    set_cell_shading(table.rows[0].cells[1], LIGHT_FILL)
    for section, page in contents:
        cells = table.add_row().cells
        set_cell_text(cells[0], section)
        set_cell_text(cells[1], page)
    doc.add_page_break()


def add_introduction(doc):
    heading(doc, "1. Introduction")
    paragraphs = [
        "The project titled AI-Powered Placement Preparation and Student Progress Tracker is developed as a database-oriented web application for managing the placement preparation activities of students. In many colleges, students prepare for interviews, coding rounds, resume screening, aptitude tests, and company applications using separate tools. The result is usually scattered data: tasks are written in notebooks, resumes are shared through messages, mock interview feedback is stored separately, and mentors do not get a continuous view of student progress. This project brings those activities into a single structured system where the data can be stored, retrieved, analyzed, and updated in an organized way.",
        "The main idea of the system is not only to create screens for students and mentors, but also to maintain a practical database model behind those screens. A student can maintain a profile, upload resume information, list target companies, track weak topics, attend mock interviews, solve coding problems, record preparation progress, and monitor placement applications. A mentor or administrator can observe the student's readiness, assign tasks, add notes, review progress, and communicate with the student. Each of these actions creates records in the database, and those records are linked through references such as userId, studentId, mentorId, problemId, and conversationId.",
        "The project is implemented using the MERN stack. MongoDB is used as the database because the application handles flexible and nested data such as resume versions, interview answers, coding test cases, notification metadata, and application timelines. Express.js and Node.js form the backend API layer. React is used for the frontend interface. Mongoose is used to define schemas, enforce validation, create indexes, and manage relationships between collections. This makes the project suitable for a DBMS report because it includes clear entity design, relationship mapping, schema implementation, indexing, validation, and CRUD operations.",
        "The application has three major user roles: student, mentor, and admin. A student is the central user of the system. The student profile stores academic details, skills, weak topics, resume metadata, placement readiness scores, and target companies. The mentor role is designed to guide students by assigning preparation plans, creating mentor notes, giving feedback, and monitoring activity. The admin role is used for broader analytics and system-level observation. Role-based access control protects the routes so that each user can access only the allowed operations.",
        "A major database requirement of this system is to keep related records connected without losing flexibility. For example, one user has one student profile, but one student profile can have many tasks, many progress logs, many mock interviews, many placement applications, many ATS reviews, and many mentor notes. Similarly, one coding problem can have many coding submissions. A conversation can contain many messages. These relationships are represented using MongoDB ObjectId references in Mongoose schemas.",
        "The system also supports AI-assisted features such as mock interview question generation, answer evaluation, resume scoring, ATS analysis, and personalized preparation guidance. From a database point of view, the important part is that generated questions, answers, scores, feedback, resume snapshots, metrics, and suggestions are saved as structured data. This allows the system to show history and trends instead of treating every AI response as temporary output.",
        "The project solves a practical problem faced by placement students. It provides a single place to plan preparation, measure progress, revise resumes, practice interviews, solve coding problems, and track applications. The database design helps the system remain consistent as the amount of data increases. Since every feature depends on well-designed collections and relationships, the project demonstrates the importance of DBMS concepts in a real application.",
    ]
    for text in paragraphs:
        para(doc, text)
    heading(doc, "1.1 Objectives", 2)
    for item in [
        "To design a database-backed placement preparation system for students, mentors, and administrators.",
        "To store student profiles, resume details, tasks, interview history, coding submissions, progress logs, and placement applications in an organized form.",
        "To use MongoDB schemas and references to model relationships between users and preparation activities.",
        "To support retrieval of dashboard analytics such as readiness score, task completion rate, consistency, and interview performance.",
        "To provide a maintainable backend structure with authentication, authorization, validation, and route-based API access.",
    ]:
        bullet(doc, item)
    heading(doc, "1.2 Scope of the Project", 2)
    para(doc, "The scope of the project covers student placement preparation from profile setup to interview practice and application tracking. It includes authentication, profile management, resume parsing and ATS review, task planning, mock interview evaluation, coding practice, mentor guidance, notifications, chat messages, progress analytics, and report generation. The project is not limited to storing static student details; it continuously records preparation activities and uses them to produce useful insights.")
    doc.add_page_break()
    heading(doc, "1.3 Need for the System", 2)
    for text in [
        "Placement preparation usually contains many small but important activities. A student may solve coding questions on one platform, prepare aptitude separately, update resumes in files, discuss doubts with mentors through chat, and track applications manually. When this information is scattered, it becomes difficult to know whether the student is actually improving. The proposed system solves this by storing each activity in a connected database.",
        "From the DBMS perspective, the project is useful because it shows how academic and preparation data can be structured. The database does not simply store one row per student. It stores a timeline of student growth through progress logs, interview scores, resume versions, task completion, coding submissions, and application status. These records can later be used to generate reports and analytics.",
        "The project also reduces duplication. Instead of storing mentor details again and again in every task or note, the system stores mentor references using ObjectId fields. Instead of copying student details into every collection, the records point back to StudentProfile. This keeps the database easier to update and reduces inconsistency.",
    ]:
        para(doc, text)
    heading(doc, "1.4 Users of the System", 2)
    add_table(doc, ["User Type", "Role in the System"], [
        ("Student", "Maintains profile, uploads resume, tracks tasks, attempts mock interviews, solves coding problems, records progress, and tracks applications."),
        ("Mentor", "Reviews student preparation, assigns tasks, writes mentor notes, adds feedback, and monitors student activity."),
        ("Admin", "Views broader analytics, manages system-level information, and monitors platform usage."),
    ])
    doc.add_page_break()


def add_schema_section(doc):
    heading(doc, "2. Database-Schema Diagram")
    para(doc, "The database schema diagram gives a collection-level view of how the application data is arranged in MongoDB. Since MongoDB is a document database, the design uses separate collections for large independent entities and embedded subdocuments for smaller details that naturally belong to a parent document. For example, interview answers are embedded inside a mock interview record, while a coding submission is stored as a separate collection because it has its own lifecycle and references both the student and the coding problem.")
    doc.add_picture(str(SCHEMA_IMG), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Figure 2.1: Database schema diagram showing major MongoDB collections and references.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "The central collections are User and StudentProfile. The User collection manages authentication, role, active status, and login-related information. StudentProfile stores academic and placement-specific details. Most preparation records are connected to StudentProfile using studentId, while authentication and ownership are connected using userId. This separation keeps login data independent from academic and placement data.")
    heading(doc, "2.1 Main Collections", 2)
    rows = [
        ("User", "Stores login identity, encrypted password, role, and account status."),
        ("StudentProfile", "Stores department, semester, target role, target companies, skills, weak topics, resume metadata, readiness scores, and mentor link."),
        ("Task", "Stores preparation tasks created by the student, mentor, or AI plan."),
        ("ProgressLog", "Stores daily or weekly study progress, completed topics, scores, and notes."),
        ("MockInterview", "Stores interview type, company, questions, answers, scores, delivery metrics, and feedback."),
        ("ATSReview", "Stores resume text, job description, matched keywords, missing keywords, formatting issues, and ATS metrics."),
        ("PlacementApplication", "Stores company applications, role details, round status, timeline events, and next actions."),
        ("CodingProblem", "Stores coding questions, difficulty, tags, test cases, supported languages, and starter code."),
        ("CodingSubmission", "Stores submitted code, verdict, score, runtime, and testcase result details."),
        ("Conversation and Message", "Store mentor-student communication and message history."),
        ("Notification and Activity", "Store alerts and activity feed events for better traceability."),
    ]
    add_table(doc, ["Collection", "Purpose"], rows)
    heading(doc, "2.2 Schema Design Notes", 2)
    for text in [
        "ObjectId references are used where one document must point to another document, such as Task.userId, Task.studentId, MockInterview.studentId, CodingSubmission.problemId, and Message.conversationId.",
        "Embedded arrays are used where the child data is meaningful mainly inside the parent document. Examples include MockInterview.answers, PlacementApplication.rounds, PlacementApplication.timeline, CodingProblem.testCases, and StudentProfile.resumeVersions.",
        "Indexes are added on frequently searched or filtered fields such as userId, studentId, role, status, createdAt, difficulty, category, company, and lastMessageAt.",
        "Validation is handled through Mongoose enum values, required fields, min/max limits, and schema-level defaults.",
    ]:
        bullet(doc, text)
    doc.add_page_break()
    heading(doc, "2.3 Collection-wise Schema Summary", 2)
    schema_rows = []
    for name, fields in entities.items():
        schema_rows.append((name, ", ".join(fields[:9])))
    add_table(doc, ["Collection", "Important Fields"], schema_rows)
    para(doc, "The schema is designed to balance normalization and document flexibility. MongoDB does not force a relational structure, but the project still follows DBMS design discipline by keeping repeated concepts in separate collections and linking them through references. This makes the database easier to query, update, and extend.")
    doc.add_page_break()
    heading(doc, "2.4 Important Field Explanation", 2)
    add_table(doc, ["Field", "Used In", "Explanation"], [
        ("userId", "Most student-owned records", "Connects a record to the authenticated user account."),
        ("studentId", "Tasks, interviews, progress, applications", "Connects activity records to a StudentProfile document."),
        ("mentorId", "StudentProfile, MentorNote, PreparationPlan", "Represents the mentor responsible for guiding a student."),
        ("status", "Task, Application, Notification, Interview", "Represents the current state of a record and supports filtering."),
        ("createdAt / updatedAt", "All timestamped models", "Maintains history and allows recent activity sorting."),
        ("metrics", "Analytics, ATSReview, MockInterview", "Stores calculated scores and performance indicators."),
        ("timeline", "PlacementApplication", "Stores chronological events in an application process."),
        ("answers", "MockInterview", "Stores question-wise answer, score, and feedback details."),
    ])
    para(doc, "The repeated use of userId and studentId is intentional. userId identifies the login account, while studentId identifies the placement profile. This is helpful because not every user is a student; mentors and admins also exist in the User collection. Student-specific preparation records therefore use studentId wherever the record belongs to an actual student profile.")
    doc.add_page_break()


def add_er_section(doc):
    heading(doc, "3. ER Diagram")
    para(doc, "The ER diagram represents the logical relationships between the main entities in the project. Although the implementation is done using MongoDB collections, ER modeling is still useful because it explains how the data objects are connected from a database design point of view.")
    doc.add_picture(str(ER_IMG), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Figure 3.1: ER diagram showing cardinality among major entities.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading(doc, "3.1 Relationship Description", 2)
    relationship_rows = [
        ("User - StudentProfile", "One-to-One", "A student user has one detailed student profile."),
        ("StudentProfile - Task", "One-to-Many", "One student can have many preparation tasks."),
        ("StudentProfile - ProgressLog", "One-to-Many", "A student can record many progress entries over time."),
        ("StudentProfile - MockInterview", "One-to-Many", "A student can attend many mock interviews."),
        ("StudentProfile - ATSReview", "One-to-Many", "A student can store multiple resume review versions."),
        ("StudentProfile - PlacementApplication", "One-to-Many", "A student can apply to many companies."),
        ("CodingProblem - CodingSubmission", "One-to-Many", "A single coding problem can receive many submissions."),
        ("Conversation - Message", "One-to-Many", "A conversation contains multiple messages."),
        ("User - Notification", "One-to-Many", "A user can receive many notifications."),
    ]
    add_table(doc, ["Relationship", "Cardinality", "Explanation"], relationship_rows)
    para(doc, "The design avoids storing all student activity inside a single document because the document would become too large and difficult to query. Instead, high-volume data such as interviews, progress logs, coding submissions, messages, and applications are stored separately. This gives the application better query performance and keeps each collection focused on one responsibility.")
    doc.add_page_break()
    heading(doc, "3.2 Cardinality Justification", 2)
    for text in [
        "The User to StudentProfile relationship is one-to-one for student accounts. A student should not have multiple active placement profiles because dashboard scores, resume details, and mentor mapping would become confusing. The StudentProfile document therefore acts as the main academic and placement identity of the student.",
        "The StudentProfile to Task relationship is one-to-many. A student may receive many tasks over time. Some tasks may be self-created, some may be assigned by a mentor, and some may come from an AI plan. Keeping tasks as separate records allows filtering by pending, completed, overdue, priority, and category.",
        "The StudentProfile to MockInterview relationship is also one-to-many. Every mock interview is a separate attempt with its own questions, answers, scores, and suggestions. Storing interviews separately allows the system to compare old and new performance instead of overwriting the previous attempt.",
        "The Conversation to Message relationship is one-to-many. This is a common design decision for chat modules because a conversation can grow continuously. Keeping messages separate makes it easier to load recent messages and update read status without rewriting a large conversation document.",
    ]:
        para(doc, text)
    doc.add_page_break()


def add_requirements(doc):
    heading(doc, "4. Hardware and Software Requirement")
    para(doc, "The project is a web-based database application. It can run on a normal development laptop during implementation and can be deployed using cloud services for production. The backend and frontend are separated, so the hardware and software requirements are divided into development requirements and deployment requirements.")
    heading(doc, "4.1 Hardware Requirements", 2)
    add_table(doc, ["Component", "Minimum Requirement", "Recommended Requirement"], [
        ("Processor", "Dual-core processor", "Intel i5 / Ryzen 5 or above"),
        ("RAM", "4 GB", "8 GB or above"),
        ("Storage", "2 GB free space", "5 GB or above for node_modules and build files"),
        ("Network", "Stable internet for MongoDB Atlas and package installation", "Broadband connection"),
        ("Display", "1366 x 768", "Full HD display for dashboard testing"),
    ])
    heading(doc, "4.2 Software Requirements", 2)
    add_table(doc, ["Software", "Purpose"], [
        ("Windows / Linux / macOS", "Operating system for development and testing."),
        ("Node.js", "Runs the Express backend and frontend build tools."),
        ("MongoDB Atlas", "Cloud database used for storing application data."),
        ("Express.js", "Backend API framework."),
        ("React.js and Vite", "Frontend user interface and development server."),
        ("Mongoose", "Schema modeling, validation, indexes, and MongoDB interaction."),
        ("Axios", "Frontend HTTP client for API communication."),
        ("JWT and bcryptjs", "Authentication token generation and password hashing."),
        ("PDFKit and document parsers", "Report generation and resume text extraction."),
        ("Visual Studio Code", "Code editor used for development."),
        ("Browser", "Used to access the application UI."),
    ])
    heading(doc, "4.3 Functional Requirements", 2)
    for item in [
        "Users must be able to register and log in securely.",
        "Students must be able to create and update their profile.",
        "The system must store resume data, parsed skills, and resume review history.",
        "Students must be able to create tasks and record progress logs.",
        "Mock interview questions, answers, scores, and feedback must be stored.",
        "Mentors must be able to monitor students and provide notes or feedback.",
        "The system must generate analytics using stored preparation data.",
        "Placement applications must be tracked with status, rounds, and timeline.",
    ]:
        bullet(doc, item)
    heading(doc, "4.4 Non-Functional Requirements", 2)
    for item in [
        "The system should protect private student data through authentication and role-based authorization.",
        "The database should support efficient searching and filtering using indexes.",
        "The API should return clear error messages and avoid exposing internal details.",
        "The frontend should remain responsive and usable on standard laptop displays.",
        "The database design should allow future features without changing the entire structure.",
    ]:
        bullet(doc, item)
    doc.add_page_break()
    heading(doc, "4.5 System Constraints and Assumptions", 2)
    for text in [
        "The application assumes that each student uses one account for placement preparation. This keeps progress tracking consistent. If a student creates multiple accounts, the preparation history would be split across different User and StudentProfile records.",
        "The project assumes that MongoDB Atlas is available through a valid connection string. If the database connection fails, the backend health check reports the database state and protected API routes cannot perform normal operations.",
        "The system assumes that the frontend communicates only with the backend API. The frontend does not contain MongoDB credentials and does not directly query the database. This separation is important for security and maintainability.",
        "AI-related features are designed with fallback output. If an API key is not configured, the system can still return sample guidance or scores. This allows the database and application flow to be tested even when external AI access is not available.",
    ]:
        para(doc, text)
    doc.add_page_break()


def add_design_implementation(doc):
    heading(doc, "5. Database Design and Implementation")
    para(doc, "Database design is the most important part of this project because almost every feature depends on stored data. The design was prepared by identifying the main actors, the data generated by each actor, and the relationships between those data items. Since the project uses MongoDB, the design follows document database principles while still maintaining logical relationships using ObjectId references.")
    heading(doc, "5.1 Database Selection", 2)
    para(doc, "MongoDB was selected because the project includes semi-structured data. Resume details, interview answers, coding testcase results, message attachments, notification metadata, and placement timelines do not always have a fixed simple tabular shape. MongoDB allows these values to be stored as nested objects and arrays. At the same time, Mongoose schemas provide enough structure to keep the data reliable.")
    heading(doc, "5.2 Mongoose Implementation", 2)
    para(doc, "The backend uses Mongoose models for each major collection. Each schema defines fields, data types, required constraints, default values, enum restrictions, indexes, and references. The models are imported into controllers and services, where create, read, update, and delete operations are performed. This gives the project a clear separation between database structure and business logic.")
    add_table(doc, ["Model", "Implementation Detail"], [
        ("User", "Stores login credentials, role, and activity fields. Passwords are hashed before saving."),
        ("StudentProfile", "Stores placement profile details and references the user and mentor."),
        ("Task", "Stores tasks with priority, status, category, source, deadline, and reminder time."),
        ("MockInterview", "Stores generated questions, answers, evaluation metrics, strengths, improvements, and status."),
        ("ATSReview", "Stores resume-review output including keyword matching and score metrics."),
        ("PlacementApplication", "Stores company-wise application progress with embedded rounds and timeline events."),
        ("CodingProblem", "Stores coding challenge content, constraints, templates, and testcases."),
        ("CodingSubmission", "Stores user code, verdict, score, and per-testcase result."),
    ])
    heading(doc, "5.3 Normalization and Embedding Decisions", 2)
    para(doc, "The database uses a mixed design. Independent records that may grow over time are stored in separate collections, while closely related child data is embedded in the parent document. For example, application rounds are embedded inside PlacementApplication because they are meaningful only for that application. On the other hand, messages are stored separately from Conversation because a conversation can contain many messages and users may need to load them page by page.")
    for item in [
        "Separate collections are used for users, student profiles, tasks, interviews, progress logs, applications, messages, and coding submissions.",
        "Embedded subdocuments are used for resume versions, interview answers, application rounds, application timeline events, coding test cases, and submission cases.",
        "References are used when one collection must be connected to another, such as userId, studentId, mentorId, problemId, conversationId, senderId, and receiverId.",
        "Indexes support common dashboard and search operations, especially for user-specific and status-specific queries.",
    ]:
        bullet(doc, item)
    doc.add_page_break()
    heading(doc, "5.3.1 Collection Design Details", 2)
    para(doc, "The User collection is kept compact because it is accessed during login and authorization. It stores only identity and account-level details such as name, email, password hash, role, and activity timestamps. Password hashing is handled before saving, so the database does not store readable passwords.")
    para(doc, "The StudentProfile collection contains the student's placement identity. It stores academic details, target companies, skills, weak topics, resume metadata, parsed resume details, readiness scores, GitHub and LinkedIn links, and mentor mapping. This document is read frequently by dashboards and interview modules, so it acts as the main anchor for student preparation data.")
    para(doc, "The Task and ProgressLog collections capture day-to-day preparation. Task records show what the student has to do, while ProgressLog records what the student has actually completed. This separation is useful because pending work and completed effort are different kinds of information.")
    para(doc, "The MockInterview collection stores interview sessions. Each session can include generated questions, answers, scoring metrics, voice or visual delivery metrics, strengths, improvements, and suggestions. This design allows every interview to remain available for later review.")
    para(doc, "The PlacementApplication collection represents the student's company application journey. It contains the company name, role, status, priority, rounds, timeline, package details, notes, and next action. Embedded rounds and timeline events are suitable here because they belong to one application record.")
    doc.add_page_break()
    heading(doc, "5.4 CRUD Operations", 2)
    add_table(doc, ["Operation", "Example in Project"], [
        ("Create", "Register user, create task, upload resume review, create placement application, submit coding solution."),
        ("Read", "Load dashboard, fetch profile, list tasks, show ATS history, view interview sessions."),
        ("Update", "Update profile, mark task completed, update application status, mark notification read."),
        ("Delete / Soft Delete", "Messages use an isDeleted field, and several records can be filtered by active status."),
    ])
    heading(doc, "5.5 Authentication and Data Protection", 2)
    para(doc, "Authentication is implemented using JSON Web Tokens. Passwords are not stored directly; bcrypt is used to hash passwords before saving them in the User collection. Protected routes check the token and attach the logged-in user to the request. Role-based middleware ensures that student, mentor, and admin actions remain separated. This is important because the database stores personal information, resume content, interview performance, and preparation feedback.")
    heading(doc, "5.6 Indexing Strategy", 2)
    para(doc, "Indexes are used on fields that are frequently used for lookup, sorting, filtering, or searching. Examples include User.email, User.role, StudentProfile.userId, Task.userId, Task.status, ProgressLog.date, MockInterview.studentId, ATSReview.userId, PlacementApplication.status, CodingProblem.slug, CodingSubmission.problemId, Conversation.participants, and Notification.status. These indexes reduce the time needed to fetch dashboard data and history records.")
    add_table(doc, ["Collection", "Index / Search Field", "Reason"], [
        ("User", "email, role, createdAt", "Login lookup and role-based filtering."),
        ("Task", "userId, status, priority, createdAt", "Dashboard task filtering and sorting."),
        ("MockInterview", "studentId, type, company, status", "Interview history and company-specific practice."),
        ("ATSReview", "userId + createdAt, company + targetRole", "Resume review history and trend comparison."),
        ("PlacementApplication", "userId + status + updatedAt", "Application tracker dashboard."),
        ("CodingProblem", "slug, difficulty, category, tags", "Coding arena filtering."),
        ("Message", "conversationId + createdAt", "Chat history loading."),
        ("Notification", "userId + status + priority", "Unread notification display."),
    ])
    doc.add_page_break()
    heading(doc, "5.7 Backend API Implementation", 2)
    para(doc, "The backend exposes REST API routes under the /api path. Each route file maps a feature area to controller functions. The student routes handle profile updates, progress logs, resume parsing, and resume analysis. Interview routes handle mock interview generation and answer evaluation. Task routes handle task CRUD operations. Analytics routes calculate student and admin dashboard data. Coding routes manage coding problems and submissions.")
    heading(doc, "5.8 Frontend and Database Interaction", 2)
    para(doc, "The React frontend does not connect directly to MongoDB. It communicates with the Express API using Axios. The API client attaches the JWT token from local storage to every protected request. This keeps database credentials and business logic on the server side. The frontend displays the returned data in forms, dashboards, charts, cards, tables, and progress sections.")
    heading(doc, "5.9 Error Handling and Validation", 2)
    para(doc, "The project uses centralized error handling so that database and validation errors are returned in a consistent format. Input validation is applied using express-validator where required, and schema-level validation is applied through Mongoose. This reduces invalid records and helps maintain database consistency.")
    heading(doc, "5.10 Data Flow Example", 2)
    for text in [
        "When a student updates the profile, the frontend sends the changed fields to the student API. The backend verifies the token, checks that the user has the student role, and updates the StudentProfile collection. The updated profile is returned to the frontend and displayed immediately.",
        "When a student uploads or analyzes a resume, the backend extracts text, parses skills and projects, stores resume metadata, updates parsedResume, and keeps a resume version history. This allows the student to compare future resume improvements.",
        "When a mock interview is completed, the questions, answers, evaluation metrics, and overall score are stored in MockInterview. The score can then contribute to dashboard analytics and preparation guidance.",
        "When a task is completed, its status and completedAt field are updated. This can improve task completion rate and helps mentors understand whether the student is following the plan.",
    ]:
        para(doc, text)
    doc.add_page_break()


def add_results(doc):
    heading(doc, "6. Results")
    para(doc, "The completed project provides a working database-backed placement preparation system. The application stores student details, preparation tasks, resume data, interview history, coding attempts, mentor feedback, notifications, and application records. The result is a connected data system where students and mentors can observe preparation progress instead of handling each activity separately.")
    heading(doc, "6.1 Major Outputs", 2)
    for item in [
        "Student profile page stores department, semester, target role, target companies, skills, weak topics, GitHub, LinkedIn, and resume metadata.",
        "Resume upload and analysis stores parsed resume text, extracted skills, projects, resume versions, and ATS review results.",
        "Mock interview module stores generated questions, answers, scoring metrics, feedback, strengths, improvements, and suggestions.",
        "Task module stores preparation tasks with deadline, category, priority, status, tags, and completion time.",
        "Progress log module stores study minutes, completed topics, scores, consistency, and notes.",
        "Application tracker stores company, role, source, status, priority, rounds, timeline, next action, and interview dates.",
        "Coding arena stores coding problems, supported languages, submissions, verdicts, scores, and testcase results.",
        "Dashboard analytics summarizes readiness, task completion, average interview score, consistency, and active usage.",
    ]:
        bullet(doc, item)
    heading(doc, "6.2 Sample Result Table", 2)
    add_table(doc, ["Feature", "Database Collection Used", "Observed Result"], [
        ("Profile Management", "StudentProfile", "Student placement metadata is saved and loaded correctly."),
        ("Resume Analysis", "StudentProfile, ATSReview", "Resume text and score history are available for review."),
        ("Mock Interview", "MockInterview", "Questions, answers, scores, and feedback are stored session-wise."),
        ("Task Planning", "Task", "Tasks can be filtered by status, priority, category, and deadline."),
        ("Coding Practice", "CodingProblem, CodingSubmission", "Problems and submissions are stored separately and linked by problemId."),
        ("Mentor Guidance", "MentorNote, Feedback, PreparationPlan", "Mentor actions remain connected to the selected student."),
        ("Communication", "Conversation, Message", "Messages are grouped by conversation and sender/receiver references."),
        ("Notifications", "Notification", "Unread and read alerts can be managed using status."),
    ])
    heading(doc, "6.3 Database Outcome", 2)
    para(doc, "The database design supports both detailed records and summarized dashboard views. The system can show current status and also preserve history. For example, a student can see the latest readiness score, while the database still stores earlier progress logs, older resume versions, past interview sessions, and previous application updates. This makes the project useful not only as a CRUD application but also as a progress tracking system.")
    doc.add_page_break()
    heading(doc, "6.4 Testing Observations", 2)
    for text in [
        "The profile workflow verifies that student data can be created, updated, and retrieved through the API. This confirms the connection between User and StudentProfile records.",
        "The task workflow verifies status changes such as pending, in-progress, completed, and overdue. This confirms that enum validation and filtering work correctly.",
        "The resume workflow verifies that resume content can be saved as metadata, parsed text, extracted skills, and review history. This confirms that nested resume fields and resumeVersions are useful in practice.",
        "The interview workflow verifies that question lists and answer evaluations can be stored in one interview session document. This confirms that embedded answer subdocuments are appropriate for the use case.",
        "The application tracker verifies that company-wise application status and timeline events can be maintained without mixing them with unrelated student details.",
    ]:
        para(doc, text)
    heading(doc, "6.5 Benefits of the Implemented Database", 2)
    for item in [
        "Student preparation data is stored in one connected system.",
        "Mentors can review history instead of relying only on verbal updates.",
        "Dashboard analytics can be generated from real database records.",
        "Resume, interview, coding, and application data remain available for future comparison.",
        "The schema can be extended for more companies, more interview rounds, or more analytics.",
    ]:
        bullet(doc, item)
    doc.add_page_break()


def add_conclusion(doc):
    heading(doc, "7. Conclusion")
    paragraphs = [
        "The AI-Powered Placement Preparation and Student Progress Tracker successfully demonstrates the use of DBMS concepts in a practical web application. The project is not limited to a simple student table; it contains a complete database design with multiple connected collections, embedded subdocuments, indexes, validation rules, authentication, role-based access, and API-driven data operations.",
        "The database design supports the real workflow of placement preparation. Students can maintain profiles, upload resume details, practice interviews, solve coding problems, manage tasks, record progress, and track applications. Mentors can guide students with notes, feedback, tasks, and preparation plans. Administrators can observe broader activity and analytics. These features are possible because the database is organized around clear entities and relationships.",
        "MongoDB is suitable for this project because the data is flexible and activity-based. Mongoose gives structure to the database while still allowing nested objects such as interview answers, resume versions, application rounds, timeline events, and testcase results. This combination gives the project both flexibility and reliability.",
        "Overall, the project shows how a well-designed database can turn scattered placement preparation activities into a useful tracking and decision-support system. Future improvements can include stronger reporting, advanced analytics, more detailed mentor dashboards, company-specific preparation plans, and improved data visualization. The existing schema provides a strong base for those extensions."
    ]
    for text in paragraphs:
        para(doc, text)
    heading(doc, "7.1 Future Enhancements", 2)
    for item in [
        "Add more detailed company-wise preparation roadmaps and connect them with tasks and progress logs.",
        "Introduce stronger analytics using historical trends from interviews, coding submissions, and resume scores.",
        "Add export options for student progress reports and mentor review summaries.",
        "Improve the notification module with scheduled reminders and calendar integration.",
        "Add more database backup and audit features for production usage.",
    ]:
        bullet(doc, item)
    doc.add_page_break()


def add_references(doc):
    heading(doc, "8. References")
    refs = [
        "MongoDB Documentation. MongoDB Manual: Documents, Collections, Indexes, and Schema Design.",
        "Mongoose Documentation. Schemas, Models, Validation, Middleware, and Query Indexing.",
        "Express.js Documentation. Routing, Middleware, Error Handling, and REST API Development.",
        "React Documentation. Components, State Management, Effects, and Frontend Rendering.",
        "Node.js Documentation. JavaScript Runtime, Modules, and Server-side Application Development.",
        "JSON Web Token Introduction. JWT-based Authentication and Authorization Concepts.",
        "bcryptjs Package Documentation. Password Hashing for Web Applications.",
        "Chart.js Documentation. Dashboard Charts and Data Visualization.",
        "MongoDB Atlas Documentation. Cloud Database Deployment and Connection String Configuration.",
        "Project Source Code. AI-Powered Placement Preparation and Student Progress Tracker, local MERN implementation."
    ]
    for idx, ref in enumerate(refs, 1):
        numbered(doc, ref)
    para(doc, "The report content is written based on the local project implementation, especially the Mongoose models, backend routes, controller responsibilities, and README description.")
    doc.add_page_break()
    heading(doc, "8.1 Project Files Referred", 2)
    para(doc, "The following local project files were used as the basis for the database explanation. They are included here to make the report traceable to the actual implementation rather than to a generic placement-management example.")
    add_table(doc, ["Project File", "Reason for Reference"], [
        ("backend/models/User.js", "Used for authentication entity fields, password hashing, and user roles."),
        ("backend/models/StudentProfile.js", "Used for student academic details, resume data, readiness scores, and mentor mapping."),
        ("backend/models/Task.js", "Used for preparation task schema, categories, priority, and status."),
        ("backend/models/MockInterview.js", "Used for interview session design, answer subdocuments, and scoring metrics."),
        ("backend/models/ATSReview.js", "Used for resume review schema, keyword matching, and ATS score metrics."),
        ("backend/models/PlacementApplication.js", "Used for company application tracking, rounds, and timeline design."),
        ("backend/models/CodingProblem.js", "Used for coding challenge structure and testcase storage."),
        ("backend/models/CodingSubmission.js", "Used for submission verdicts, testcase results, and problem references."),
        ("backend/models/Conversation.js and Message.js", "Used for mentor-student communication relationship design."),
        ("README.md", "Used for project overview, architecture, technology stack, and setup description."),
    ])


def build():
    create_schema_diagram()
    create_er_diagram()
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_toc(doc)
    add_introduction(doc)
    add_schema_section(doc)
    add_er_section(doc)
    add_requirements(doc)
    add_design_implementation(doc)
    add_results(doc)
    add_conclusion(doc)
    add_references(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
